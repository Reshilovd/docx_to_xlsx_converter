from io import BytesIO

from docx import Document
from openpyxl import Workbook


def convert_docx_to_xlsx(docx_bytes: bytes) -> bytes:
    """Convert tables from DOCX into XLSX workbook."""
    document = Document(BytesIO(docx_bytes))
    workbook = Workbook()
    default_sheet = workbook.active

    has_tables = False
    for index, table in enumerate(document.tables, start=1):
        has_tables = True
        sheet_name = f"table_{index}"

        if index == 1:
            sheet = default_sheet
            sheet.title = sheet_name
        else:
            sheet = workbook.create_sheet(title=sheet_name)

        for row_index, row in enumerate(table.rows, start=1):
            for col_index, cell in enumerate(row.cells, start=1):
                sheet.cell(row=row_index, column=col_index, value=cell.text.strip())

    if not has_tables:
        default_sheet.title = "content"
        full_text = "\n".join(
            paragraph.text.strip()
            for paragraph in document.paragraphs
            if paragraph.text.strip()
        )
        default_sheet.cell(row=1, column=1, value=full_text or "Document is empty.")

    output = BytesIO()
    workbook.save(output)
    output.seek(0)
    return output.read()
